import sys

from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

path = r"D:\LanguageCenter\BaoCao_WebPrograming_working.docx"
doc = Document(path)

print("PARAGRAPHS")
for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.replace("\t", " | ").strip()
    if text:
        print(f"P{index:03d} [{paragraph.style.name}] {text}")

print("\nTABLES")
for table_index, table in enumerate(doc.tables):
    print(f"TABLE {table_index}: {len(table.rows)}x{len(table.columns)}")
    for row_index, row in enumerate(table.rows):
        values = [" ".join(cell.text.split()) for cell in row.cells]
        print(f"  R{row_index:02d}: {values}")

print("\nSECTIONS", len(doc.sections))
print("INLINE_SHAPES", len(doc.inline_shapes))
