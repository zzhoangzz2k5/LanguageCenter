import math
import os
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORK = Path(r"D:\LanguageCenter")
INPUT = WORK / "BaoCao_WebPrograming_working.docx"
OUTPUT = WORK / "BaoCao_WebPrograming_Part1_Part2.docx"
ERD_PATH = WORK / "language-center-erd.png"
USE_CASE_PATH = WORK / "language-center-use-case.png"

BLUE = "1F4E78"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"
INK = "1F1F1F"


def pil_font(size, bold=False):
    filename = "arialbd.ttf" if bold else "arial.ttf"
    return ImageFont.truetype(str(Path(r"C:\Windows\Fonts") / filename), size)


def rounded_box(draw, box, title, fields, fill, outline=(45, 75, 110), title_fill=None):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=4)
    header_height = 56
    if title_fill:
        draw.rounded_rectangle(
            (x1, y1, x2, y1 + header_height + 12),
            radius=18,
            fill=title_fill,
            outline=outline,
            width=4,
        )
        draw.rectangle((x1 + 2, y1 + header_height, x2 - 2, y1 + header_height + 15), fill=title_fill)
    draw.text((x1 + 18, y1 + 10), title, font=pil_font(34, True), fill=(255, 255, 255) if title_fill else (20, 45, 70))
    y = y1 + header_height + 12
    for field in fields:
        draw.text((x1 + 18, y), field, font=pil_font(27), fill=(30, 30, 30))
        y += 35


def center(box):
    x1, y1, x2, y2 = box
    return ((x1 + x2) // 2, (y1 + y2) // 2)


def connect(draw, box_a, box_b, label_a="1", label_b="N"):
    ax, ay = center(box_a)
    bx, by = center(box_b)
    dx = bx - ax
    dy = by - ay
    if abs(dx) > abs(dy):
        start = (box_a[2] if dx > 0 else box_a[0], ay)
        end = (box_b[0] if dx > 0 else box_b[2], by)
    else:
        start = (ax, box_a[3] if dy > 0 else box_a[1])
        end = (bx, box_b[1] if dy > 0 else box_b[3])
    draw.line((start, end), fill=(70, 90, 115), width=4)
    draw.ellipse((start[0] - 5, start[1] - 5, start[0] + 5, start[1] + 5), fill=(70, 90, 115))
    draw.ellipse((end[0] - 5, end[1] - 5, end[0] + 5, end[1] + 5), fill=(70, 90, 115))
    draw.text((start[0] + 7, start[1] - 25), label_a, font=pil_font(18, True), fill=(20, 50, 85))
    draw.text((end[0] - 24, end[1] - 25), label_b, font=pil_font(18, True), fill=(20, 50, 85))


def make_erd():
    image = Image.new("RGB", (2500, 1700), "white")
    draw = ImageDraw.Draw(image)
    draw.text((90, 45), "LANGUAGE CENTER MANAGEMENT SYSTEM - ERD", font=pil_font(44, True), fill=(31, 78, 121))
    draw.text((90, 102), "PK = Primary Key   FK = Foreign Key", font=pil_font(23), fill=(80, 80, 80))

    boxes = {
        "UserAccount": (80, 200, 570, 560),
        "Student": (700, 170, 1180, 440),
        "Teacher": (700, 500, 1180, 770),
        "Program": (1300, 150, 1830, 510),
        "Class": (1300, 550, 1830, 970),
        "Registration": (1930, 550, 2440, 880),
        "Payment": (1930, 970, 2440, 1230),
        "PlacementTest": (700, 1030, 1210, 1420),
        "Consultation": (80, 1040, 600, 1370),
        "ClassSchedule": (1300, 1050, 1830, 1370),
    }
    entity_fields = {
        "UserAccount": ["PK UserId", "FullName", "Email", "PasswordHash", "Role", "IsActive", "CreatedDate", "Photo"],
        "Student": ["PK StudentId", "FK UserId", "Phone", "Address", "Avatar"],
        "Teacher": ["PK TeacherId", "FK UserId", "Specialty", "ExperienceYears", "TeacherImage"],
        "Program": ["PK ProgramId", "ProgramName", "LevelName", "DurationMonths", "Fee", "Description", "ProgramImage", "ProgramStatus"],
        "Class": ["PK ClassId", "FK ProgramId", "FK TeacherId", "ClassName", "Room", "StartDate", "EndDate", "Capacity", "Status"],
        "Registration": ["PK RegistrationId", "FK StudentId", "FK ClassId", "RegistrationDate", "Status"],
        "Payment": ["PK PaymentId", "FK RegistrationId", "Amount", "PaymentDate", "PaymentMethod", "PaymentStatus"],
        "PlacementTest": ["PK PlacementTestId", "FK StudentId", "TestDate", "TestTime", "SuggestedLevel", "ResultScore", "Status"],
        "Consultation": ["PK ConsultationId", "FK StudentId", "Question", "ContactInfo", "RequestStatus", "CreatedDate"],
        "ClassSchedule": ["PK ScheduleId", "FK ClassId", "DayOfWeek", "StartTime", "EndTime"],
    }
    fills = {
        "UserAccount": (235, 243, 250),
        "Student": (235, 248, 239),
        "Teacher": (255, 245, 225),
        "Program": (240, 236, 250),
        "Class": (232, 244, 252),
        "Registration": (252, 238, 238),
        "Payment": (245, 245, 225),
        "PlacementTest": (236, 246, 246),
        "Consultation": (246, 238, 248),
        "ClassSchedule": (238, 242, 248),
    }
    for name, box in boxes.items():
        rounded_box(draw, box, name, entity_fields[name], fills[name], title_fill=(31, 78, 121))

    relationships = [
        ("UserAccount", "Student", "1", "0..1"),
        ("UserAccount", "Teacher", "1", "0..1"),
        ("Program", "Class", "1", "N"),
        ("Teacher", "Class", "1", "N"),
        ("Student", "Registration", "1", "N"),
        ("Class", "Registration", "1", "N"),
        ("Registration", "Payment", "1", "N"),
        ("Student", "PlacementTest", "1", "N"),
        ("Student", "Consultation", "1", "N"),
        ("Class", "ClassSchedule", "1", "N"),
    ]
    for a, b, la, lb in relationships:
        connect(draw, boxes[a], boxes[b], la, lb)

    draw.text(
        (90, 1535),
        "Core process: User account -> Student/Teacher profile -> Program -> Class -> Registration -> Payment",
        font=pil_font(25, True),
        fill=(31, 78, 121),
    )
    draw.text(
        (90, 1580),
        "Supporting processes: Placement tests, consultations, and class schedules.",
        font=pil_font(23),
        fill=(65, 65, 65),
    )
    image.save(ERD_PATH, quality=95)


def actor(draw, x, y, name):
    draw.ellipse((x - 28, y - 100, x + 28, y - 44), outline=(31, 78, 121), width=5)
    draw.line((x, y - 44, x, y + 45), fill=(31, 78, 121), width=5)
    draw.line((x - 52, y - 10, x + 52, y - 10), fill=(31, 78, 121), width=5)
    draw.line((x, y + 45, x - 48, y + 105), fill=(31, 78, 121), width=5)
    draw.line((x, y + 45, x + 48, y + 105), fill=(31, 78, 121), width=5)
    text_box = draw.textbbox((0, 0), name, font=pil_font(27, True))
    draw.text((x - (text_box[2] - text_box[0]) / 2, y + 125), name, font=pil_font(27, True), fill=(31, 78, 121))


def use_case(draw, box, text, fill):
    draw.ellipse(box, fill=fill, outline=(55, 80, 110), width=4)
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    total_h = len(lines) * 34
    y = (y1 + y2 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=pil_font(28, True))
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, font=pil_font(28, True), fill=(25, 35, 45))
        y += 34


def make_use_case():
    image = Image.new("RGB", (2500, 1600), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 40), "LANGUAGE CENTER MANAGEMENT SYSTEM - USE CASE DIAGRAM", font=pil_font(42, True), fill=(31, 78, 121))
    draw.rounded_rectangle((420, 140, 2080, 1480), radius=25, outline=(31, 78, 121), width=5)
    draw.text((450, 160), "System Boundary", font=pil_font(26, True), fill=(31, 78, 121))

    actor_positions = {
        "Guest": (190, 370),
        "Student": (190, 1050),
        "Teacher": (2310, 430),
        "Admin": (2310, 1120),
    }
    for name, (x, y) in actor_positions.items():
        actor(draw, x, y, name)

    cases = [
        ("Browse programs\nand classes", (520, 230, 980, 340), (232, 244, 252), "Guest"),
        ("Register / Login", (1110, 230, 1510, 340), (232, 244, 252), "Guest"),
        ("View teachers\nand home content", (1610, 230, 2010, 340), (232, 244, 252), "Guest"),
        ("Manage profile", (500, 520, 860, 625), (235, 248, 239), "Student"),
        ("Register class", (930, 520, 1290, 625), (235, 248, 239), "Student"),
        ("View classes\nand payments", (1360, 520, 1760, 625), (235, 248, 239), "Student"),
        ("Book placement test", (700, 690, 1110, 795), (235, 248, 239), "Student"),
        ("Submit consultation", (1220, 690, 1630, 795), (235, 248, 239), "Student"),
        ("Teacher dashboard", (1530, 390, 1990, 495), (255, 245, 225), "Teacher"),
        ("View teaching classes\nand students", (1530, 850, 1990, 970), (255, 245, 225), "Teacher"),
        ("Upload / delete\nclass materials", (1530, 1030, 1990, 1150), (255, 245, 225), "Teacher"),
        ("View tests and\nconsultations", (1530, 1210, 1990, 1330), (255, 245, 225), "Teacher"),
        ("Dashboard statistics", (500, 900, 900, 1005), (246, 238, 248), "Admin"),
        ("Manage users", (500, 1080, 900, 1185), (246, 238, 248), "Admin"),
        ("Manage programs\nand classes", (960, 900, 1390, 1020), (246, 238, 248), "Admin"),
        ("Manage registrations\nand payments", (960, 1080, 1390, 1200), (246, 238, 248), "Admin"),
        ("Manage placement\ntests", (960, 1260, 1390, 1380), (246, 238, 248), "Admin"),
    ]
    for text, box, fill, role in cases:
        use_case(draw, box, text, fill)
        ax, ay = actor_positions[role]
        target_x = box[0] if ax < 1250 else box[2]
        target_y = (box[1] + box[3]) // 2
        start_x = ax + 65 if ax < 1250 else ax - 65
        draw.line((start_x, ay - 10, target_x, target_y), fill=(110, 120, 135), width=3)

    draw.text((540, 1430), "Authentication and role checks protect Student, Teacher, and Admin functions.", font=pil_font(23), fill=(70, 70, 70))
    image.save(USE_CASE_PATH, quality=95)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_run(run, size=12, bold=False, color=INK, italic=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, after=5, before=0, line_spacing=1.15, alignment=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        paragraph.alignment = alignment


def style_table(table, widths=None, font_size=9):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B7B7B7")
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if widths:
                cell.width = Inches(widths[col_index])
            if row_index == 0:
                shade_cell(cell, BLUE)
            elif row_index % 2 == 0:
                shade_cell(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph, after=0, line_spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT)
                for run in paragraph.runs:
                    style_run(run, size=font_size, bold=(row_index == 0), color=WHITE if row_index == 0 else INK)
    set_repeat_table_header(table.rows[0])


def delete_between(start_paragraph, end_paragraph):
    start = start_paragraph._p
    end = end_paragraph._p
    node = start.getnext()
    while node is not None and node is not end:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node


def move_before(anchor_paragraph, element):
    anchor_paragraph._p.addprevious(element)


def ensure_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Symbol")
    fonts.set(qn("w:hAnsi"), "Symbol")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_paragraph_before(doc, anchor, text="", size=12, bold=False, color=INK, before=0, after=5, alignment=None, keep_with_next=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold, color=color)
    format_paragraph(paragraph, before=before, after=after, alignment=alignment)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    move_before(anchor, paragraph._p)
    return paragraph


def add_bullets_before(doc, anchor, items, level=0):
    num_id = getattr(doc, "_report_bullet_num_id", None)
    if num_id is None:
        num_id = ensure_bullet_numbering(doc)
        doc._report_bullet_num_id = num_id
    for item in items:
        paragraph = doc.add_paragraph()
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        p_pr.append(num_pr)
        run = paragraph.add_run(item)
        style_run(run, size=12)
        format_paragraph(paragraph, after=4)
        move_before(anchor, paragraph._p)


def add_picture_before(doc, anchor, path, width, caption):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    format_paragraph(paragraph, after=3)
    move_before(anchor, paragraph._p)
    caption_p = add_paragraph_before(
        doc,
        anchor,
        caption,
        size=10,
        color="666666",
        after=8,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for run in caption_p.runs:
        run.italic = True


def add_table_before(doc, anchor, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    style_table(table, widths, font_size)
    move_before(anchor, table._tbl)
    spacer = add_paragraph_before(doc, anchor, "", size=4, after=4)
    return table


def add_page_break_before(doc, anchor):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    move_before(anchor, paragraph._p)


def main():
    make_erd()
    make_use_case()

    doc = Document(INPUT)
    part1 = next(p for p in doc.paragraphs if p.text.strip().startswith("Part 1."))
    part3 = next(p for p in doc.paragraphs if p.text.strip().startswith("Part 3."))
    topic = next(p for p in doc.paragraphs if p.text.strip().startswith("TOPIC 2."))
    topic.paragraph_format.page_break_before = True
    delete_between(part1, part3)

    # Part 1
    add_paragraph_before(doc, part3, "1.1 Problem Statement", size=14, bold=True, color=BLUE, before=8, after=6, keep_with_next=True)
    add_paragraph_before(
        doc,
        part3,
        "Language centers commonly manage learners, teachers, courses, class schedules, registrations, placement tests, and tuition payments through separate spreadsheets, paper forms, messaging applications, and manual confirmation. This fragmented process makes information difficult to search, creates duplicate records, delays registration approval, and increases the risk of incorrect class capacity, payment status, or test results.",
    )
    add_paragraph_before(
        doc,
        part3,
        "The problem is important because a language center serves several groups at the same time. Learners need accurate course information and a simple registration process; teachers need timely class lists and teaching materials; administrators need reliable operational data. When these groups work with disconnected information, the center spends more time correcting records and less time supporting learning quality.",
    )
    add_paragraph_before(
        doc,
        part3,
        "The Language Center Management Website provides one shared web-based system. Guests can explore programs, classes, and teachers. Registered students can update profiles, register for classes, review payments, book placement tests, and submit consultation requests. Teachers can monitor assigned classes, view students, and manage class materials. Administrators can approve accounts and manage programs, classes, teachers, students, registrations, payments, and placement-test results.",
    )
    add_bullets_before(
        doc,
        part3,
        [
            "For learners: faster access to program details, class availability, registration status, payment information, and placement-test schedules.",
            "For teachers: a consolidated teaching dashboard, class roster, schedule overview, and document-upload area.",
            "For the center: centralized records, role-based workflows, operational statistics, reduced manual work, and more consistent service.",
        ],
    )

    add_paragraph_before(doc, part3, "1.2 Topic Selection", size=14, bold=True, color=BLUE, before=10, after=6, keep_with_next=True)
    add_paragraph_before(
        doc,
        part3,
        "The team selected this topic because language-center operations combine public information, user authentication, role-based authorization, database relationships, transactional workflows, file uploads, search, filtering, pagination, and dashboard statistics. It is therefore suitable for demonstrating the major learning outcomes of the Web Programming course while solving a recognizable real-world management problem.",
    )
    add_paragraph_before(doc, part3, "Target users", size=12, bold=True, color=BLUE, before=4, after=4, keep_with_next=True)
    add_bullets_before(
        doc,
        part3,
        [
            "Guest: explores the center, programs, classes, and teacher information; creates an account and logs in.",
            "Student: maintains a personal profile, registers for classes, follows registration/payment status, books a placement test, and requests consultation.",
            "Teacher: views assigned classes and students, checks placement-test and consultation information, and manages learning materials.",
            "Admin: approves user requests and manages the center's complete operational dataset and dashboard.",
        ],
    )

    add_paragraph_before(doc, part3, "1.3 Survey of Existing Websites", size=14, bold=True, color=BLUE, before=10, after=6, keep_with_next=True)
    add_paragraph_before(
        doc,
        part3,
        "The survey was conducted on 9 June 2026 using the public home pages of VUS and Wall Street English Vietnam. The comparison focuses on information architecture and learner-facing functions rather than internal systems that are not publicly visible.",
        size=11,
        color="555555",
    )
    add_paragraph_before(doc, part3, "Website 1: VUS", size=13, bold=True, color=BLUE, before=8, after=4, keep_with_next=True)
    add_paragraph_before(doc, part3, "Official website: https://vus.edu.vn/", size=11, color="555555", after=5)
    add_picture_before(doc, part3, WORK / "vus-home.png", 6.4, "Figure 1. VUS home page captured on 9 June 2026.")
    add_table_before(
        doc,
        part3,
        ["Aspect", "Analysis"],
        [
            ("Advantages", "Strong visual branding; clear course groups by age and learning need; prominent trial-test, support, and consultation actions; rich learning resources."),
            ("Disadvantages", "The home page contains many promotional blocks and competing calls to action; key operational information such as class capacity, registration status, and payment progress is not presented as one simple workflow."),
            ("Features adopted", "Program categories, visual program cards, teacher/center credibility, placement-test entry point, consultation call to action, and responsive navigation."),
        ],
        [1.35, 5.0],
        font_size=10,
    )

    add_paragraph_before(doc, part3, "Website 2: Wall Street English Vietnam", size=13, bold=True, color=BLUE, before=8, after=4, keep_with_next=True)
    add_paragraph_before(doc, part3, "Official website: https://www.wallstreetenglish.edu.vn/", size=11, color="555555", after=5)
    add_picture_before(doc, part3, WORK / "wall-street-english-home.png", 6.4, "Figure 2. Wall Street English Vietnam home page captured on 9 June 2026.")
    add_table_before(
        doc,
        part3,
        ["Aspect", "Analysis"],
        [
            ("Advantages", "Clean hierarchy; clear value proposition; course cards for different learner goals; visible level-test and student-login links; strong consultation call to action."),
            ("Disadvantages", "Course pricing and detailed class availability are not immediately visible; the public journey is consultation-oriented, so users cannot see a complete registration/payment workflow before contacting the center."),
            ("Features adopted", "Simple hero message, goal-based program discovery, student login, level-test access, clear program detail pages, and focused calls to action."),
        ],
        [1.35, 5.0],
        font_size=10,
    )
    add_paragraph_before(doc, part3, "Proposed feature set after the survey", size=12, bold=True, color=BLUE, before=8, after=4, keep_with_next=True)
    add_bullets_before(
        doc,
        part3,
        [
            "Public home page with banner, statistics, featured programs, new classes, teachers, and clear Login/Register links.",
            "Program list with search, level filter, pagination, fee, duration, image, and detail page.",
            "Integrated account approval and role assignment instead of a purely promotional lead form.",
            "Student self-service for class registration, payment tracking, placement-test booking, and consultation requests.",
            "Teacher and Admin workspaces so the public website and internal operations share one database.",
        ],
    )

    # Part 2
    part2_heading = add_paragraph_before(doc, part3, "Part 2. System Design (15 Points)", size=15, bold=True, color=INK, before=0, after=8, keep_with_next=True)
    part2_heading.paragraph_format.page_break_before = True
    add_paragraph_before(doc, part3, "2.1 ERD Diagram", size=14, bold=True, color=BLUE, before=6, after=6, keep_with_next=True)
    add_picture_before(doc, part3, ERD_PATH, 6.4, "Figure 3. Entity Relationship Diagram of the Language Center Management System.")
    add_paragraph_before(
        doc,
        part3,
        "The database contains ten implemented entities. UserAccount stores authentication and role information. Student and Teacher extend a user account with role-specific profile data. Program and Class describe the academic offering. Registration links a student to a class, while Payment records tuition transactions for that registration. PlacementTest and Consultation support learner services, and ClassSchedule represents recurring teaching times.",
    )
    add_table_before(
        doc,
        part3,
        ["Table", "Primary key", "Purpose"],
        [
            ("UserAccount", "UserId", "Authentication, identity, role, active status, and profile photo."),
            ("Student", "StudentId", "Student-specific phone, address, and avatar; references UserAccount."),
            ("Teacher", "TeacherId", "Teacher specialty, experience, and image; references UserAccount."),
            ("Program", "ProgramId", "Course name, level, duration, fee, description, image, and status."),
            ("Class", "ClassId", "Program offering assigned to a teacher, room, dates, capacity, and status."),
            ("ClassSchedule", "ScheduleId", "Day and start/end time for a class."),
            ("Registration", "RegistrationId", "Connects a student with a class and stores registration status."),
            ("Payment", "PaymentId", "Amount, date, method, and status for a registration."),
            ("PlacementTest", "PlacementTestId", "Test appointment, suggested level, score, and status."),
            ("Consultation", "ConsultationId", "Student question, contact information, request status, and date."),
        ],
        [1.25, 1.15, 4.0],
        font_size=9,
    )
    add_paragraph_before(doc, part3, "Key relationships", size=12, bold=True, color=BLUE, before=8, after=4, keep_with_next=True)
    add_bullets_before(
        doc,
        part3,
        [
            "One UserAccount may have one Student profile or one Teacher profile, depending on the approved role.",
            "One Program has many Classes; each Class belongs to one Program.",
            "One Teacher may teach many Classes; each Class is assigned to zero or one Teacher during planning.",
            "Student and Class have a many-to-many relationship resolved by Registration.",
            "One Registration may have multiple Payment records to support transaction history or installments.",
            "One Student may create many PlacementTest and Consultation records.",
            "One Class may contain many ClassSchedule records.",
        ],
    )

    add_paragraph_before(doc, part3, "2.2 Use Case Diagram", size=14, bold=True, color=BLUE, before=10, after=6, keep_with_next=True)
    add_picture_before(doc, part3, USE_CASE_PATH, 6.4, "Figure 4. Use Case Diagram with Guest, Student, Teacher, and Admin actors.")
    add_table_before(
        doc,
        part3,
        ["Role", "Main use cases"],
        [
            ("Guest", "View home page, browse/search/filter programs, view program details and classes, view teachers, register, and log in."),
            ("Student", "View/update profile, change password, upload avatar, register for a class, view classes and payments, book placement test, and submit consultation request."),
            ("Teacher", "View dashboard, assigned classes and student lists, upload/delete class materials, view placement-test results and consultation feedback, and upload profile photo."),
            ("Admin", "View dashboard statistics; approve accounts; manage students, teachers, programs, classes, registrations, payments, and placement tests."),
        ],
        [1.25, 5.15],
        font_size=10,
    )
    add_paragraph_before(
        doc,
        part3,
        "Authorization is based on Session values (UserId and Role). Public actions are available to guests, while Student, Teacher, and Admin actions verify the current role before accessing role-specific data.",
    )

    add_paragraph_before(doc, part3, "2.3 Sample Data", size=14, bold=True, color=BLUE, before=10, after=6, keep_with_next=True)
    add_paragraph_before(
        doc,
        part3,
        "The following dataset is designed for demonstration and testing. It meets the minimum requirement: 5 programs, 10 classes, 10 students, 3 teachers, 20 registrations, and 10 placement-test records.",
    )

    programs = [
        ("P01", "General English Foundation", "A1", "3", "3,500,000", "Active"),
        ("P02", "English Communication", "A2-B1", "4", "4,800,000", "Active"),
        ("P03", "IELTS Preparation", "B1-B2", "6", "8,500,000", "Active"),
        ("P04", "TOEIC Intensive", "A2-B2", "4", "5,900,000", "Active"),
        ("P05", "Business English", "B1-C1", "5", "7,200,000", "Active"),
    ]
    add_paragraph_before(doc, part3, "Programs (5 records)", size=12, bold=True, color=BLUE, before=6, after=4, keep_with_next=True)
    add_table_before(doc, part3, ["ID", "Program name", "Level", "Months", "Fee (VND)", "Status"], programs, [0.45, 2.35, 0.75, 0.65, 1.15, 0.75], 8)

    teachers = [
        ("T01", "Nguyen Minh Anh", "IELTS / Academic English", "8"),
        ("T02", "Tran Hoang Nam", "Communication / TOEIC", "6"),
        ("T03", "Le Thu Ha", "Business / General English", "10"),
    ]
    add_paragraph_before(doc, part3, "Teachers (3 records)", size=12, bold=True, color=BLUE, before=6, after=4, keep_with_next=True)
    add_table_before(doc, part3, ["ID", "Full name", "Specialty", "Experience (years)"], teachers, [0.55, 1.8, 2.8, 1.3], 9)

    classes = [
        ("C01", "GEF-A1-Morning", "P01", "T03", "A101", "01/07/2026", "30/09/2026", "20", "Open"),
        ("C02", "GEF-A1-Evening", "P01", "T03", "A102", "06/07/2026", "05/10/2026", "20", "Open"),
        ("C03", "COM-A2-Weekend", "P02", "T02", "B201", "04/07/2026", "31/10/2026", "18", "Open"),
        ("C04", "COM-B1-Evening", "P02", "T02", "B202", "13/07/2026", "13/11/2026", "18", "Open"),
        ("C05", "IELTS-5.5", "P03", "T01", "C301", "01/08/2026", "31/01/2027", "16", "Open"),
        ("C06", "IELTS-6.5", "P03", "T01", "C302", "08/08/2026", "07/02/2027", "16", "Planned"),
        ("C07", "TOEIC-550", "P04", "T02", "B203", "20/07/2026", "20/11/2026", "22", "Open"),
        ("C08", "TOEIC-750", "P04", "T02", "B204", "27/07/2026", "27/11/2026", "20", "Planned"),
        ("C09", "Business-B1", "P05", "T03", "D401", "03/08/2026", "03/01/2027", "18", "Open"),
        ("C10", "Business-B2", "P05", "T03", "D402", "10/08/2026", "10/01/2027", "18", "Planned"),
    ]
    add_paragraph_before(doc, part3, "Classes (10 records)", size=12, bold=True, color=BLUE, before=6, after=4, keep_with_next=True)
    add_table_before(doc, part3, ["ID", "Class", "Program", "Teacher", "Room", "Start", "End", "Cap.", "Status"], classes, [0.35, 1.4, 0.55, 0.55, 0.55, 0.8, 0.8, 0.45, 0.75], 7)

    students = [
        ("S01", "Pham Gia Bao", "bao.pham@example.com", "0901000001", "District 1"),
        ("S02", "Nguyen Ngoc Lan", "lan.nguyen@example.com", "0901000002", "District 3"),
        ("S03", "Tran Minh Khoa", "khoa.tran@example.com", "0901000003", "Binh Thanh"),
        ("S04", "Le Hoang Yen", "yen.le@example.com", "0901000004", "District 7"),
        ("S05", "Vo Thanh Tung", "tung.vo@example.com", "0901000005", "Thu Duc"),
        ("S06", "Do My Linh", "linh.do@example.com", "0901000006", "Go Vap"),
        ("S07", "Bui Quoc Huy", "huy.bui@example.com", "0901000007", "Tan Binh"),
        ("S08", "Hoang Thu Trang", "trang.hoang@example.com", "0901000008", "Phu Nhuan"),
        ("S09", "Dang Anh Tuan", "tuan.dang@example.com", "0901000009", "District 10"),
        ("S10", "Mai Khanh Vy", "vy.mai@example.com", "0901000010", "District 5"),
    ]
    add_paragraph_before(doc, part3, "Students (10 records)", size=12, bold=True, color=BLUE, before=6, after=4, keep_with_next=True)
    add_table_before(doc, part3, ["ID", "Full name", "Email", "Phone", "Address"], students, [0.45, 1.45, 2.35, 1.15, 1.0], 8)

    registrations = []
    pairs = [
        ("R01", "S01", "C05", "Confirmed"), ("R02", "S01", "C03", "Pending"),
        ("R03", "S02", "C01", "Confirmed"), ("R04", "S02", "C07", "Confirmed"),
        ("R05", "S03", "C04", "Pending"), ("R06", "S03", "C05", "Confirmed"),
        ("R07", "S04", "C02", "Confirmed"), ("R08", "S04", "C09", "Pending"),
        ("R09", "S05", "C03", "Confirmed"), ("R10", "S05", "C07", "Pending"),
        ("R11", "S06", "C01", "Confirmed"), ("R12", "S06", "C04", "Confirmed"),
        ("R13", "S07", "C05", "Pending"), ("R14", "S07", "C09", "Confirmed"),
        ("R15", "S08", "C02", "Confirmed"), ("R16", "S08", "C03", "Confirmed"),
        ("R17", "S09", "C07", "Confirmed"), ("R18", "S09", "C09", "Pending"),
        ("R19", "S10", "C04", "Confirmed"), ("R20", "S10", "C05", "Pending"),
    ]
    for index, (rid, sid, cid, status) in enumerate(pairs, start=1):
        registrations.append((rid, sid, cid, f"{index:02d}/06/2026", status))
    add_paragraph_before(doc, part3, "Registrations (20 records)", size=12, bold=True, color=BLUE, before=6, after=4, keep_with_next=True)
    add_table_before(doc, part3, ["ID", "Student", "Class", "Registration date", "Status"], registrations, [0.55, 0.8, 0.8, 1.7, 1.2], 8)

    tests = [
        ("PT01", "S01", "15/06/2026", "08:00", "B1", "62", "Completed"),
        ("PT02", "S02", "15/06/2026", "09:00", "A1", "35", "Completed"),
        ("PT03", "S03", "16/06/2026", "08:00", "B2", "78", "Completed"),
        ("PT04", "S04", "16/06/2026", "09:00", "A2", "49", "Completed"),
        ("PT05", "S05", "17/06/2026", "08:00", "B1", "65", "Completed"),
        ("PT06", "S06", "17/06/2026", "09:00", "A2", "52", "Completed"),
        ("PT07", "S07", "18/06/2026", "08:00", "B1", "-", "Pending"),
        ("PT08", "S08", "18/06/2026", "09:00", "A2", "-", "Pending"),
        ("PT09", "S09", "19/06/2026", "08:00", "B2", "-", "Pending"),
        ("PT10", "S10", "19/06/2026", "09:00", "B1", "-", "Pending"),
    ]
    add_paragraph_before(doc, part3, "Placement tests (10 records)", size=12, bold=True, color=BLUE, before=6, after=4, keep_with_next=True)
    add_table_before(doc, part3, ["ID", "Student", "Date", "Time", "Level", "Score", "Status"], tests, [0.6, 0.75, 1.05, 0.75, 0.75, 0.7, 1.05], 8)

    add_paragraph_before(
        doc,
        part3,
        "Data integrity rules used during testing: email addresses are unique; a student cannot register for the same class twice; class and program status values control availability; payment and registration statuses are updated only by authorized actions; foreign keys must reference existing records.",
        size=11,
        color="555555",
        before=6,
        after=8,
    )

    # Keep Part 3 visually separate.
    part3.paragraph_format.page_break_before = True

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
